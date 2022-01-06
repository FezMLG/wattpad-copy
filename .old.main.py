import os
import shutil
from urllib.parse import urljoin

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_BREAK
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.common.keys import Keys
from bcolors import bcolors
from os.path import basename
from ebook import makeEpub
from ebooklib import epub


headers = {
    "User-Agent": 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) '
                  'Chrome/90.0.4430.212 Safari/537.36'}

print(bcolors.LightMagenta + "Podaj link do pierwszego rozdziału: " + bcolors.ResetAll)
startURL = input()
startURL.strip()
URL = startURL

print(bcolors.Green + "Przygotowywanie" + bcolors.ResetAll)
try:
    browser = webdriver.Firefox()
    browser.get(URL)
    html_source = browser.page_source
    browser.quit()
except:
    print(bcolors.Red + "Złe URL!" + bcolors.ResetAll)
    browser.quit()
    exit()

soup = BeautifulSoup(html_source, 'html.parser')
print(bcolors.Green + "Zakończono przygotowania" + bcolors.ResetAll)

book_title = soup.find(class_="title h5").get_text().strip()
print(bcolors.Green + "Tytuł książki:", book_title + bcolors.ResetAll)

document = Document()
print(bcolors.Green + "Utworzono nowy plik docx" + bcolors.ResetAll)

document.add_heading(book_title, 0)


def getChapterTitle(soup):
    title_div = soup.find(class_="row part-header")
    title = title_div.find("h1", class_="h2").get_text().strip()
    return title


def getText(soup):
    ret = ""
    contents = soup.find_all("div", class_="page")
    for content in contents:
        out = content.find(class_="panel").get_text()
        ret += out
    return ret.strip()


def writeTextDoc(title, text, document):
    p = document.add_heading(title, level=1)
    run = p.add_run()
    run.add_break()

    p = document.add_paragraph(text)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def generateChapters(title, text):
    c1 = epub.EpubHtml(title=title)
    c1.content = '<html><head></head><body><h1>'+title+'</h1><p>'+text+'</p></body></html>'
    return c1


def nextURL(soup):
    try:
        next_url_is = soup.find(class_="on-navigate next-part-link")["href"]
    except:
        next_url_is = ""
    return next_url_is


def get_images(soup, url):
    images = [img for img in soup.findAll('img', class_="cover")]
    print(str(len(images)) + " images found.")
    print('Downloading cover photo.')
    image_links = [each.get('src') for each in images]
    for each in image_links:
        print(f"{each}: ")
        try:
            filename = each.strip().split('/')[-1].strip()
            src = urljoin(url, each)
            print('Getting: ' + filename)
            response = requests.get(src, stream=True)
            # delay to avoid corrupted previews
            # time.sleep(1)
            with open(filename, 'wb') as out_file:
                shutil.copyfileobj(response.raw, out_file)
        except:
            print('An error occured. Continuing.')
    print('Done.')
    return filename


while not URL == "":

    for span in soup.find_all("span", {'class': 'comment-marker'}):
        span.decompose()

    for br in soup.find_all("br"):
        br.replace_with("\n")

    title = getChapterTitle(soup)
    text = getText(soup)
    print(bcolors.Green + "Dodawanie:", title + bcolors.ResetAll)
    writeTextDoc(title, text, document)
    URL = nextURL(soup)
    if URL == "":
        break
    print(bcolors.LightGreen + "Dodano:", title + bcolors.ResetAll)
    page = requests.get(URL, headers=headers)
    soup = BeautifulSoup(page.content, 'html.parser')

file_name = "".join(x for x in book_title if x.isalnum())

print(bcolors.LightMagenta + "Nazwa pliku (" + file_name + "):" + bcolors.ResetAll)
file_name_user = input()

if not file_name_user == "":
    file_name = file_name_user

document.save(file_name + '.docx')
print(bcolors.Green + "Książkę zapisano pod nazwą", file_name + bcolors.ResetAll)
